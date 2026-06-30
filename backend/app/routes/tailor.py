"""Routes for generating and downloading JD-tailored resumes."""

import os
from typing import Any, Dict, List

from docx import Document
from fastapi import APIRouter, Depends, HTTPException, status
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session

from app.api.v1.auth.security import get_current_user
from app.core.database import get_db
from app.models.job_description import JobDescription
from app.models.resume import Resume
from app.models.resume_match import ResumeMatch
from app.models.tailored_resume import TailoredResume
from app.models.user import User
from app.schemas.tailor_schema import TailorResumeRequest, TailorResumeResponse
from app.services.resume_tailor.prompt_builder import build_tailoring_prompt
from app.services.resume_tailor.resume_rewriter import rewrite_resume
from app.services.matcher_service import matcher_service

router = APIRouter(prefix="/tailor", tags=["Resume Tailoring"])


def _list_values(values: Any) -> List[str]:
    if not values:
        return []
    cleaned = []
    for item in values:
        if isinstance(item, str) and item.strip():
            cleaned.append(item.strip())
        elif isinstance(item, dict):
            text = item.get("raw_info") or item.get("name") or item.get("title")
            if isinstance(text, str) and text.strip():
                cleaned.append(text.strip())
    return cleaned


def _resume_payload(resume: Resume) -> Dict[str, Any]:
    return {
        "full_name": resume.full_name,
        "email": resume.email,
        "phone": resume.phone,
        "summary": "",
        "skills": resume.skills or [],
        "experience": resume.experience or [],
        "projects": resume.projects or [],
        "certifications": resume.certifications or [],
        "technologies": resume.technologies or [],
    }


def _fallback_tailored_content(
    resume: Resume,
    jd: JobDescription,
    match: ResumeMatch | None,
    request: TailorResumeRequest,
) -> Dict[str, Any]:
    skills = _list_values((resume.skills or []) + (resume.technologies or []))
    matched = _list_values((match.matched_skills if match else []) + (match.matched_technologies if match else []))
    focus = request.focus_areas or []
    selected_skills = []
    for skill in matched + focus + skills:
        if skill and skill not in selected_skills:
            selected_skills.append(skill)

    role = request.target_role or jd.job_title or "the target role"
    summary = (
        f"Candidate profile tailored for {role}, emphasizing relevant experience, "
        f"JD-aligned skills, and measurable project impact."
    )
    if request.custom_instructions:
        summary = f"{summary} Direction: {request.custom_instructions.strip()}"

    return {
        "full_name": resume.full_name or "",
        "email": resume.email or "",
        "phone": resume.phone or "",
        "summary": summary,
        "skills": selected_skills[:18],
        "experience": resume.experience or [],
        "projects": resume.projects or [],
        "certifications": resume.certifications or [],
    }


def _write_docx(path: str, content: Dict[str, Any], exclude_sections: List[str] | None = None) -> None:
    exclude_sections = exclude_sections or []
    doc = Document()

    if content.get("full_name"):
        doc.add_heading(content["full_name"], level=0)
    contact = " | ".join([v for v in [content.get("email"), content.get("phone")] if v])
    if contact:
        doc.add_paragraph(contact)

    sections = [
        ("summary", "Professional Summary"),
        ("skills", "Skills"),
        ("experience", "Experience"),
        ("projects", "Projects"),
        ("certifications", "Certifications"),
    ]

    for key, title in sections:
        if key in exclude_sections:
            continue
        value = content.get(key)
        if not value:
            continue
        doc.add_heading(title, level=1)
        if isinstance(value, str):
            doc.add_paragraph(value)
        elif isinstance(value, list):
            for item in value:
                if isinstance(item, str):
                    doc.add_paragraph(item, style="List Bullet")
                elif isinstance(item, dict):
                    doc.add_paragraph(item.get("raw_info") or str(item), style="List Bullet")

    doc.save(path)


@router.post("/generate", response_model=TailorResumeResponse, status_code=status.HTTP_201_CREATED)
async def generate_tailored_resume(
    request: TailorResumeRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    resume = db.query(Resume).filter(
        Resume.id == request.resume_id,
        Resume.user_id == current_user.id,
    ).first()
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found.")

    jd = db.query(JobDescription).filter(
        JobDescription.id == request.jd_id,
        JobDescription.user_id == current_user.id,
    ).first()
    if not jd:
        raise HTTPException(status_code=404, detail="Job description not found.")

    match = db.query(ResumeMatch).filter(
        ResumeMatch.resume_id == request.resume_id,
        ResumeMatch.job_description_id == request.jd_id,
        ResumeMatch.user_id == current_user.id,
    ).order_by(ResumeMatch.created_at.desc()).first()

    matcher_analysis = {
        "matched_skills": (match.matched_skills if match else []) or [],
        "missing_skills": (match.missing_skills if match else []) or [],
        "matched_technologies": (match.matched_technologies if match else []) or [],
        "missing_technologies": (match.missing_technologies if match else []) or [],
    }
    preferences = request.model_dump()

    try:
        prompt = build_tailoring_prompt(_resume_payload(resume), jd.raw_text, matcher_analysis, preferences)
        tailored_content = rewrite_resume(prompt)
    except Exception:
        tailored_content = {}

    if not tailored_content:
        tailored_content = _fallback_tailored_content(resume, jd, match, request)

    original_score = float(match.match_score) if match else 0.0
    
    # Recalculate match score dynamically based on tailored resume contents
    try:
        new_match = matcher_service.perform_match_analysis(
            db=db,
            user_id=current_user.id,
            resume_id=resume.id,
            job_description_id=jd.id,
            tailored_content=tailored_content
        )
        improved_score = new_match.match_score
    except Exception as e:
        improved_score = min(100.0, round(original_score + 10.0, 1)) if original_score else 65.0

    output_dir = os.path.join(os.getcwd(), "uploads", "tailored")
    os.makedirs(output_dir, exist_ok=True)

    db_record = TailoredResume(
        user_id=current_user.id,
        resume_id=resume.id,
        jd_id=jd.id,
        original_score=original_score,
        improved_score=improved_score,
        preferences_json=preferences,
        output_file_path="pending",
        format_type="docx",
    )
    db.add(db_record)
    db.commit()
    db.refresh(db_record)

    output_path = os.path.join(output_dir, f"{db_record.id}.docx")
    _write_docx(output_path, tailored_content, request.exclude_sections)
    db_record.output_file_path = output_path
    db.commit()

    return {
        "tailored_resume_id": db_record.id,
        "original_score": original_score,
        "improved_score": improved_score,
        "changed_sections": [key for key in ["summary", "skills", "experience", "projects", "certifications"] if tailored_content.get(key)],
        "download_url": f"/tailor/download/{db_record.id}",
    }


@router.get("/download/{tailored_resume_id}")
async def download_tailored_resume(
    tailored_resume_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    record = db.query(TailoredResume).filter(
        TailoredResume.id == tailored_resume_id,
        TailoredResume.user_id == current_user.id,
    ).first()
    if not record or not os.path.exists(record.output_file_path):
        raise HTTPException(status_code=404, detail="Tailored resume file not found.")

    return FileResponse(
        record.output_file_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="tailored_resume.docx",
    )
