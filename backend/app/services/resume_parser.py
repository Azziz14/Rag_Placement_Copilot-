"""Service module for parsing resumes using PDF/DOCX extractors, Regex, SpaCy, and section heuristics.
"""

import os
import re
import spacy
import pdfplumber
import docx
from typing import List, Dict, Any

# Load spaCy model for NER
try:
    nlp = spacy.load("en_core_web_sm")
except Exception:
    nlp = None

# A standard list of skills and technologies to match against the text
COMMON_SKILLS = [
    "communication", "leadership", "teamwork", "problem solving",
    "project management", "agile", "scrum", "data structures", "algorithms",
    "system design", "software engineering", "databases", "cloud computing",
    "machine learning", "deep learning", "artificial intelligence", "web development",
    "unit testing", "ci/cd", "microservices", "restful api", "devops",
    "data analysis", "cybersecurity", "problem-solving", "critical thinking"
]

COMMON_TECHNOLOGIES = [
    "python", "java", "javascript", "typescript", "c", "c++", "c#", "golang", "rust",
    "ruby", "php", "swift", "html", "css", "sql", "postgresql", "mysql", "sqlite",
    "mongodb", "redis", "cassandra", "qdrant", "docker", "kubernetes", "aws", "azure",
    "gcp", "firebase", "supabase", "fastapi", "django", "flask", "spring boot",
    "react", "angular", "vue", "next.js", "node.js", "express", "pytorch",
    "tensorflow", "pandas", "numpy", "scikit-learn", "git", "linux", "html5", "css3"
]

# Section pattern detection
SECTION_PATTERNS = {
    "education": re.compile(r"^\b(education|academic background|academics|qualifications|studies)\b", re.IGNORECASE),
    "experience": re.compile(r"^\b(experience|work experience|professional experience|employment history|work history|career history)\b", re.IGNORECASE),
    "projects": re.compile(r"^\b(projects|personal projects|academic projects|key projects)\b", re.IGNORECASE),
    "certifications": re.compile(r"^\b(certifications|certifications & courses|courses|credentials|licenses)\b", re.IGNORECASE)
}


def extract_text(file_path: str) -> str:
    """Extract full text from PDF or DOCX file."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text
    elif ext in [".docx", ".doc"]:
        doc = docx.Document(file_path)
        text = []
        for para in doc.paragraphs:
            text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text.append(cell.text)
        return "\n".join(text)
    else:
        raise ValueError(f"Unsupported file extension: {ext}")


def extract_name(text: str) -> str:
    """Extract full name using SpaCy NER or heuristic fallback."""
    lines = [line.strip() for line in text.split("\n") if line.strip()]
    if not lines:
        return ""

    if nlp:
        first_chunk = "\n".join(lines[:5])
        doc = nlp(first_chunk)
        for ent in doc.ents:
            if ent.label_ == "PERSON":
                ent_text = ent.text.strip()
                if not re.search(r"(\b(email|phone|github|linkedin|http|@)\b)", ent_text, re.IGNORECASE):
                    return ent_text

    # Heuristic fallback: check first 3 lines
    for line in lines[:3]:
        if re.match(r"^[A-Za-z\s\.\-\']+$", line) and len(line.split()) >= 2:
            if not any(term in line.lower() for term in ["resume", "cv", "curriculum", "vitae"]):
                return line

    return lines[0] if lines else ""


def extract_email(text: str) -> str:
    """Extract email address using regex."""
    pattern = r"\b[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}\b"
    match = re.search(pattern, text)
    return match.group(0) if match else ""


def extract_phone(text: str) -> str:
    """Extract phone number using regex."""
    pattern = r"\b(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b"
    match = re.search(pattern, text)
    return match.group(0) if match else ""


def extract_skills(text: str) -> List[str]:
    """Extract matching skills from the text."""
    matched = []
    text_lower = text.lower()
    for skill in COMMON_SKILLS:
        pattern = r"\b" + re.escape(skill) + r"\b"
        if re.search(pattern, text_lower):
            matched.append(skill)
    return matched


def extract_technologies(text: str) -> List[str]:
    """Extract matching technologies from the text."""
    matched = []
    text_lower = text.lower()
    for tech in COMMON_TECHNOLOGIES:
        pattern = r"\b" + re.escape(tech) + r"\b"
        if re.search(pattern, text_lower):
            matched.append(tech)
    return matched


def _get_sections(text: str) -> Dict[str, List[str]]:
    """Divide the resume text into sections based on keywords/headings."""
    lines = [line.strip() for line in text.split("\n")]
    sections: Dict[str, List[str]] = {
        "education": [],
        "experience": [],
        "projects": [],
        "certifications": [],
        "unknown": []
    }
    
    current_section = "unknown"
    for line in lines:
        if not line:
            continue
            
        found_header = False
        for sec_name, pattern in SECTION_PATTERNS.items():
            if pattern.match(line) and len(line.split()) <= 4:
                current_section = sec_name
                found_header = True
                break
        
        if not found_header:
            sections[current_section].append(line)
            
    return sections


def extract_education(text: str) -> List[Dict[str, Any]]:
    """Extract education details using section heuristic."""
    sections = _get_sections(text)
    edu_lines = sections["education"]
    if not edu_lines:
        return []
    return [{"raw_info": "\n".join(edu_lines)}]


def extract_experience(text: str) -> List[Dict[str, Any]]:
    """Extract experience details using section heuristic."""
    sections = _get_sections(text)
    exp_lines = sections["experience"]
    if not exp_lines:
        return []
    return [{"raw_info": "\n".join(exp_lines)}]


def extract_projects(text: str) -> List[Dict[str, Any]]:
    """Extract projects using section heuristic."""
    sections = _get_sections(text)
    proj_lines = sections["projects"]
    if not proj_lines:
        return []
    return [{"raw_info": "\n".join(proj_lines)}]


def extract_certifications(text: str) -> List[str]:
    """Extract certifications using section heuristic."""
    sections = _get_sections(text)
    cert_lines = sections["certifications"]
    if cert_lines:
        return [line for line in cert_lines if len(line) > 5]
    return []


def parse_resume_file(file_path: str, original_filename: str) -> Dict[str, Any]:
    """Orchestrates extraction and parses resume fields."""
    extracted_text = extract_text(file_path)
    
    return {
        "full_name": extract_name(extracted_text),
        "email": extract_email(extracted_text),
        "phone": extract_phone(extracted_text),
        "skills": extract_skills(extracted_text),
        "education": extract_education(extracted_text),
        "experience": extract_experience(extracted_text),
        "projects": extract_projects(extracted_text),
        "certifications": extract_certifications(extracted_text),
        "technologies": extract_technologies(extracted_text),
        "extracted_text": extracted_text,
        "original_filename": original_filename
    }
