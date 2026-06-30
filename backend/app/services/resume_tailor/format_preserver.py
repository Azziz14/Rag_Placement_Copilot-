"""Module for extracting and preserving resume formatting, layout, section order, and typography.
Supports PDF and DOCX files.
"""

import os
import re
from typing import Dict, Any, List
import docx
import pdfplumber
import fitz  # PyMuPDF
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

# Reuse the patterns from resume_parser
SECTION_PATTERNS = {
    "education": re.compile(r"^\b(education|academic background|academics|qualifications|studies)\b", re.IGNORECASE),
    "experience": re.compile(r"^\b(experience|work experience|professional experience|employment history|work history|career history)\b", re.IGNORECASE),
    "projects": re.compile(r"^\b(projects|personal projects|academic projects|key projects)\b", re.IGNORECASE),
    "certifications": re.compile(r"^\b(certifications|certifications & courses|courses|credentials|licenses)\b", re.IGNORECASE),
    "skills": re.compile(r"^\b(skills|technical skills|skills & expertise|core competencies|technologies)\b", re.IGNORECASE),
    "summary": re.compile(r"^\b(summary|professional summary|about me|career objective|profile)\b", re.IGNORECASE)
}

def extract_format_info(file_path: str) -> Dict[str, Any]:
    """Extract format details such as section order, headings, bullet patterns,
    spacing rules, indentation, and font hierarchy.
    """
    ext = os.path.splitext(file_path)[1].lower()
    
    # Defaults
    format_info = {
        "section_order": [],
        "heading_styles": {
            "font_name": "Helvetica-Bold",
            "font_size": 14,
            "color": "#000000",
            "alignment": 0
        },
        "bullet_patterns": {"style": "•"},
        "spacing_rules": {"line_spacing": 1.15, "paragraph_spacing": 10},
        "indentation_rules": {"left_margin": 72, "right_margin": 72},
        "font_hierarchy": {
            "title": {"font_name": "Helvetica-Bold", "font_size": 22},
            "heading": {"font_name": "Helvetica-Bold", "font_size": 14},
            "subheading": {"font_name": "Helvetica-Bold", "font_size": 11},
            "body": {"font_name": "Helvetica", "font_size": 10}
        }
    }
    
    # 1. Section Order Extraction
    text = ""
    if ext == ".pdf":
        text = _extract_pdf_text_simple(file_path)
        format_info.update(_extract_pdf_formatting(file_path))
    elif ext in [".docx", ".doc"]:
        text = _extract_docx_text_simple(file_path)
        format_info.update(_extract_docx_formatting(file_path))
        
    format_info["section_order"] = _detect_section_order(text)
    return format_info

def _extract_pdf_text_simple(file_path: str) -> str:
    text = ""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            pt = page.extract_text()
            if pt:
                text += pt + "\n"
    return text

def _extract_docx_text_simple(file_path: str) -> str:
    doc = docx.Document(file_path)
    return "\n".join([p.text for p in doc.paragraphs])

def _detect_section_order(text: str) -> List[str]:
    lines = [line.strip() for line in text.split("\n") if line.strip()]
    order = []
    for line in lines:
        for sec_name, pattern in SECTION_PATTERNS.items():
            if pattern.match(line) and len(line.split()) <= 4:
                if sec_name not in order:
                    order.append(sec_name)
    return order

def _extract_pdf_formatting(file_path: str) -> Dict[str, Any]:
    formatting = {}
    try:
        doc = fitz.open(file_path)
        fonts = {}
        for page in doc:
            for block in page.get_text("dict")["blocks"]:
                if "lines" in block:
                    for line in block["lines"]:
                        for span in line["spans"]:
                            font_name = span["font"]
                            font_size = span["size"]
                            fonts[font_name] = fonts.get(font_name, 0) + len(span["text"])
        
        # Sort fonts by frequency
        sorted_fonts = sorted(fonts.items(), key=lambda x: x[1], reverse=True)
        if sorted_fonts:
            primary_font = sorted_fonts[0][0]
            # Map standard pdf fonts
            pdf_font = "Helvetica"
            if "times" in primary_font.lower():
                pdf_font = "Times-Roman"
            elif "courier" in primary_font.lower():
                pdf_font = "Courier"
            elif "arial" in primary_font.lower():
                pdf_font = "Helvetica"
                
            formatting["font_hierarchy"] = {
                "title": {"font_name": f"{pdf_font}-Bold", "font_size": 20},
                "heading": {"font_name": f"{pdf_font}-Bold", "font_size": 13},
                "subheading": {"font_name": f"{pdf_font}-Bold", "font_size": 11},
                "body": {"font_name": pdf_font, "font_size": 10}
            }
    except Exception:
        pass
    return formatting

def _extract_docx_formatting(file_path: str) -> Dict[str, Any]:
    formatting = {}
    try:
        doc = docx.Document(file_path)
        # Check standard bullet characters
        bullets = ["•", "-", "*", "▪"]
        found_bullet = "•"
        for p in doc.paragraphs:
            t = p.text.strip()
            if t:
                for b in bullets:
                    if t.startswith(b):
                        found_bullet = b
                        break
        formatting["bullet_patterns"] = {"style": found_bullet}
    except Exception:
        pass
    return formatting


def _set_paragraph_text(paragraph: docx.text.paragraph.Paragraph, new_text: str) -> None:
    """Replace the text in a DOCX paragraph while preserving the run-level formatting
    (font name, size, bold, italic, color) of the first run.

    Direct assignment to ``paragraph.text`` silently deletes all existing runs and
    creates a single unstyled run, which erases bold/italic/font-size metadata.
    This helper clears runs but copies the style of the first run to the new content.
    """
    # Capture the formatting of the first run before clearing
    first_run_font = None
    if paragraph.runs:
        src = paragraph.runs[0]
        first_run_font = {
            "bold": src.bold,
            "italic": src.italic,
            "underline": src.underline,
            "font_name": src.font.name,
            "font_size": src.font.size,
        }

    # Remove every existing run XML element from the paragraph
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)

    # Add a single new run with the captured formatting
    new_run = paragraph.add_run(new_text)
    if first_run_font:
        new_run.bold = first_run_font["bold"]
        new_run.italic = first_run_font["italic"]
        new_run.underline = first_run_font["underline"]
        if first_run_font["font_name"]:
            new_run.font.name = first_run_font["font_name"]
        if first_run_font["font_size"]:
            new_run.font.size = first_run_font["font_size"]


def rebuild_docx(original_path: str, output_path: str, tailored_content: Dict[str, Any], request_params: Dict[str, Any]) -> None:
    """Reconstruct DOCX with tailored content, maintaining all formatting."""
    doc = docx.Document(original_path)
    exclude_sections = request_params.get("exclude_sections") or []
    
    current_section = "unknown"
    
    # We will identify paragraphs that belong to sections, clear their text, and populate them.
    # Keep track of paragraphs to delete or modify.
    paragraphs_to_remove = []
    
    i = 0
    while i < len(doc.paragraphs):
        p = doc.paragraphs[i]
        text_strip = p.text.strip()
        
        # Check if heading
        found_section = None
        for sec_name, pattern in SECTION_PATTERNS.items():
            if pattern.match(text_strip) and len(text_strip.split()) <= 4:
                found_section = sec_name
                break
                
        if found_section:
            current_section = found_section
            if current_section in exclude_sections:
                # Remove the heading paragraph itself
                paragraphs_to_remove.append(p)
            i += 1
            continue
            
        if current_section in exclude_sections:
            paragraphs_to_remove.append(p)
            i += 1
            continue
            
        # If the paragraph belongs to a tailored section, we rewrite it
        if current_section in tailored_content and current_section != "unknown":
            section_data = tailored_content[current_section]
            
            # If it's a string (like summary or raw experience)
            if isinstance(section_data, str):
                # Use helper to preserve run-level formatting (bold, italic, font)
                _set_paragraph_text(p, section_data)
                current_section = "unknown"  # Reset so subsequent paragraphs are not overwritten
            elif isinstance(section_data, list):
                # It's experience / projects / skills / certifications
                if all(isinstance(x, str) for x in section_data):
                    # Write first bullet into this paragraph, preserving its run style
                    _set_paragraph_text(p, section_data[0])
                    current_p = p
                    # Clone formatting and insert additional paragraphs after current_p
                    for item in section_data[1:]:
                        new_p = doc.add_paragraph(style=p.style)
                        # Insert in XML structure right after current_p
                        current_p._p.addnext(new_p._p)
                        _set_paragraph_text(new_p, item)
                        current_p = new_p
                    current_section = "unknown"
                else:
                    # List of dicts (experience, projects) — serialize raw_info blocks
                    raw_texts = []
                    for item in section_data:
                        if isinstance(item, dict) and "raw_info" in item:
                            raw_texts.append(item["raw_info"])
                        elif isinstance(item, str):
                            raw_texts.append(item)
                    _set_paragraph_text(p, "\n".join(raw_texts))
                    current_section = "unknown"
                    
        i += 1
        
    # Remove excluded paragraphs
    for p in paragraphs_to_remove:
        p_element = p._p
        parent = p_element.getparent()
        if parent is not None:
            parent.remove(p_element)
            
    doc.save(output_path)


def rebuild_pdf(original_path: str, output_path: str, tailored_content: Dict[str, Any], format_info: Dict[str, Any], request_params: Dict[str, Any]) -> None:
    """Generate tailored PDF matching format_info hierarchy using ReportLab."""
    doc_template = SimpleDocTemplate(
        output_path,
        pagesize=letter,
        leftMargin=format_info["indentation_rules"]["left_margin"],
        rightMargin=format_info["indentation_rules"]["right_margin"],
        topMargin=54,
        bottomMargin=54
    )
    
    styles = getSampleStyleSheet()
    font_h = format_info["font_hierarchy"]
    
    # Define style objects
    title_style = ParagraphStyle(
        'ResumeTitle',
        parent=styles['Normal'],
        fontName=font_h["title"]["font_name"],
        fontSize=font_h["title"]["font_size"],
        leading=font_h["title"]["font_size"] + 4,
        alignment=1, # Center
        spaceAfter=15
    )
    
    heading_style = ParagraphStyle(
        'ResumeHeading',
        parent=styles['Normal'],
        fontName=font_h["heading"]["font_name"],
        fontSize=font_h["heading"]["font_size"],
        leading=font_h["heading"]["font_size"] + 4,
        spaceBefore=12,
        spaceAfter=6,
        keepWithNext=True
    )
    
    body_style = ParagraphStyle(
        'ResumeBody',
        parent=styles['Normal'],
        fontName=font_h["body"]["font_name"],
        fontSize=font_h["body"]["font_size"],
        leading=font_h["body"]["font_size"] + 4,
        spaceAfter=6
    )
    
    story = []
    
    # 1. Top Section - Name and Contact
    name = tailored_content.get("full_name") or ""
    if name:
        story.append(Paragraph(name, title_style))
        
    contact_parts = []
    email = tailored_content.get("email") or ""
    phone = tailored_content.get("phone") or ""
    if email:
        contact_parts.append(email)
    if phone:
        contact_parts.append(phone)
    if contact_parts:
        story.append(Paragraph(" | ".join(contact_parts), ParagraphStyle('Contact', parent=body_style, alignment=1, spaceAfter=15)))
        
    # 2. Iterate through sections in order
    exclude_sections = request_params.get("exclude_sections") or []
    section_order = format_info.get("section_order") or ["summary", "experience", "projects", "skills", "education", "certifications"]
    
    for sec in section_order:
        if sec in exclude_sections:
            continue
            
        content = tailored_content.get(sec)
        if not content:
            continue
            
        # Add heading
        story.append(Paragraph(sec.upper(), heading_style))
        
        # Add content
        if isinstance(content, str):
            for para in content.split("\n"):
                if para.strip():
                    story.append(Paragraph(para, body_style))
        elif isinstance(content, list):
            for item in content:
                if isinstance(item, str):
                    bullet_style = format_info["bullet_patterns"]["style"]
                    story.append(Paragraph(f"{bullet_style} {item}", body_style))
                elif isinstance(item, dict) and "raw_info" in item:
                    for para in item["raw_info"].split("\n"):
                        if para.strip():
                            story.append(Paragraph(para, body_style))
                            
    doc_template.build(story)
